import os
from datetime import datetime
from PIL import ImageGrab, ImageTk
from tkinter import Tk, Label, Entry, Button, filedialog, messagebox, PhotoImage
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

#pip install pillow python-docx

class ScreenshotApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Auto Screenshot Documenter")
        self.root.geometry("1024x720")
        icon = PhotoImage(file="./img/folder.png")  # Load the image correctly
        root.iconphoto(True, icon)  # Set the window icon


        # Variables for paths
        self.screenshot_dir = None
        self.doc_path = None
        self.current_screenshot = None
        self.screenshot_preview = None
        self.previous_image = None

        # UI Elements for path selection
        Label(root, text="Select Screenshot Directory:").pack(pady=5)
        self.screenshot_entry = Entry(root, width=50)
        self.screenshot_entry.pack(pady=5)
        Button(root, text="Browse", command=self.set_screenshot_dir).pack()

        Label(root, text="Select Document Directory:").pack(pady=5)
        self.doc_entry = Entry(root, width=50)
        self.doc_entry.pack(pady=5)
        Button(root, text="Browse", command=self.set_doc_dir).pack()

        Label(root, text="Enter Exercise number:").pack(pady=5)
        self.exercise_entry  = Entry(root, width=50)
        self.exercise_entry.pack(pady=5)

        Label(root, text="Enter Main Title:").pack(pady=5)
        self.title_entry  = Entry(root, width=50)
        self.title_entry.pack(pady=5)

        Label(root, text="Enter Name and Surname:").pack(pady=5)
        self.name_entry  = Entry(root, width=50)
        self.name_entry.pack(pady=5)

        Label(root, text="Enter your Class:").pack(pady=5)
        self.class_entry  = Entry(root, width=50)
        self.class_entry.pack(pady=5)

        # Buttons
        Button(root, text="OK", command=self.start_monitoring, 
               fg="black", bg="lime", font=("Arial", 12), relief="raised").pack(pady=10)

        Label()
        # Place the button at bottom-right with margin
        Button(root, text="CLOSE", command=root.quit).place(relx=1.0, rely=1.0, x=-10, y=-10, anchor="se")


        # Placeholder for screenshot preview
        self.image_label = Label(root)  
        self.image_label.pack(pady=10)

        self.prompt_entry = Entry(root, width=50)
        self.prompt_entry.pack(pady=5)
        self.prompt_entry.insert(0, "Here enter the screenshot description")
        self.prompt_entry.pack_forget()  # Hide initially

        self.save_button = Button(root, text="SAVE", command=self.save_screenshot,
                                fg="black", bg="lime", font=("Arial", 12, "bold"), relief="raised")
        self.save_button.pack_forget()

        self.discard_button = Button(root, text="DISCARD", command=self.discard_screenshot,
                                fg="black", bg="red", font=("Arial", 12, "bold"), relief="raised" )
        self.discard_button.pack_forget()

    def set_screenshot_dir(self):
        """Set the directory where screenshots will be saved."""
        directory = filedialog.askdirectory()
        if directory:
            self.screenshot_dir = directory
            self.screenshot_entry.delete(0, "end")
            self.screenshot_entry.insert(0, directory)

    def set_doc_dir(self):
        """Set the directory where the document will be saved."""
        directory = filedialog.askdirectory()
        if directory:
            self.doc_path = os.path.join(directory, "screenshots_document.docx")
            self.doc_entry.delete(0, "end")
            self.doc_entry.insert(0, self.doc_path)

    def start_monitoring(self):
        """Start monitoring clipboard for screenshots."""
        if not self.screenshot_dir or not self.doc_path:
            messagebox.showwarning("Warning", "Please select valid directories first!")
            return

        # Hide initial setup UI
        for widget in self.root.winfo_children():
            widget.pack_forget()
        self.screenshot_label = Label(root, text="No new screenshots taken", fg="gray")
        self.screenshot_label.pack(pady=50)
        self.image_label.pack()

        # Start checking clipboard
        self.check_clipboard()

    def check_clipboard(self):
        """Continuously monitor the clipboard for new screenshots."""
        try:
            image = ImageGrab.grabclipboard()
            if isinstance(image, ImageGrab.Image.Image) and image != self.previous_image:
                self.display_screenshot(image)
                self.previous_image = image
            else:
                self.root.after(1000, self.check_clipboard)
        except KeyboardInterrupt:
            print("\nStopping clipboard monitoring. Goodbye!")

    def display_screenshot(self, image):
        """Display the screenshot in the UI for user approval."""
        self.current_screenshot = image.copy()  # Keep the full-quality image for saving

    # Get the image size
        width, height = image.size
        
        # Define a maximum width and height for the display
        max_width, max_height = 400, 300

        # Calculate the aspect ratio
        aspect_ratio = width / height

        # Resize the image while maintaining the aspect ratio
        if width > max_width or height > max_height:
            if aspect_ratio > 1:  # Landscape image
                new_width = max_width
                new_height = int(new_width / aspect_ratio)
            else:  # Portrait or square image
                new_height = max_height
                new_width = int(new_height * aspect_ratio)
            
            resized_image = image.resize((new_width, new_height))
        else:
            resized_image = image

        # Convert the image to a Tkinter-compatible photo image
        self.screenshot_preview = ImageTk.PhotoImage(resized_image)

        # Update UI
        self.image_label.config(image=self.screenshot_preview)
        self.image_label.pack(pady=5)
        self.screenshot_label.config(text="New screenshot detected!")
        self.prompt_entry.pack(pady=5)
        self.save_button.pack(pady=5)
        self.discard_button.pack(pady=5)

    def save_screenshot(self):
        """Saves the screenshot and user-provided description to the document."""
        if self.current_screenshot is None:
            return

        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        filename = os.path.join(self.screenshot_dir, f"screenshot_{timestamp}.png")
        description = self.prompt_entry.get().strip() or "Here enter the screenshot description"

        try:
            # Save the screenshot as a PNG
            self.current_screenshot.save(filename, "PNG")

            # Load or create the document
            if os.path.exists(self.doc_path):
                doc = Document(self.doc_path)
            else:
                doc = Document()
                self.create_document_header(doc)

            # Add description and image
            self.add_styled_text(doc, description, font_size=12)
            doc.add_picture(filename, width=Inches(6))

            # Save document
            doc.save(self.doc_path)

            print(f"Screenshot saved: {filename}")
            messagebox.showinfo("Success", "Screenshot saved to document!")

            self.reset_screen()
        except Exception as e:
            messagebox.showerror("Error", f"Could not save screenshot: {e}")

    def discard_screenshot(self):
        """Discards the current screenshot."""
        self.current_screenshot = None
        messagebox.showinfo("Discarded", "Screenshot discarded.")
        self.reset_screen()

    def reset_screen(self):
        """Resets the UI back to 'no new screenshots' state."""
        self.screenshot_label.config(text="No new screenshots taken")
        self.prompt_entry.pack_forget()
        self.save_button.pack_forget()
        self.discard_button.pack_forget()
        self.image_label.config(image="")  # Clear image preview
        self.image_label.pack_forget()  # Hide image label
        self.current_screenshot = None

        # Clear the clipboard to prevent the same screenshot from being detected again
        self.root.clipboard_clear()
        
        # Delay clipboard checking to ensure the UI resets properly
        self.root.after(1500, self.check_clipboard)

    def create_document_header(self, doc):
        """Creates a header for the document on the first page."""
        self.add_styled_text(doc, self.exercise_entry.get().strip() or "Exercise Number", bold=True, font_size=16, align="center")
        self.add_styled_text(doc, self.title_entry.get().strip() or "Title", bold=True, font_size=14, align="center")
        self.add_styled_text(doc, self.name_entry.get().strip() or "Name and Surname", font_size=12, align="center")
        self.add_styled_text(doc, self.class_entry.get().strip() or "Class", font_size=12, align="center")
        self.add_styled_text(doc, f"Datum: {datetime.now().strftime('%d.%m.%Y')}", font_size=12, align="center")
        doc.add_page_break()  # Start screenshots on the next page

    def add_styled_text(self, doc, text, bold=False, font_size=14, align="left"):
        """Adds a styled paragraph to the document."""
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold
        run.font.size = Pt(font_size)

        if align == "center":
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif align == "right":
            para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


if __name__ == "__main__":
    root = Tk()
    app = ScreenshotApp(root)
    root.mainloop()
